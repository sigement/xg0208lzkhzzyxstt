import re
import os
import jieba
import matplotlib.pyplot as plt
from docx import Document
from wordcloud import WordCloud
from collections import Counter
import numpy as np
from PIL import Image
import platform
STOPWORDS = set([
    '的', '了', '是', '我', '你', '他', '她', '它', '们', '在', '和', '有', '就', '不', '也', '都', '而',
    '及', '与', '之', '于', '为', '以', '可', '将', '对', '或', '一个', '没有', '我们', '你们', '他们',
    '这里', '那里', '什么', '怎么', '哪里', '如何', '如果', '因为', '所以', '但是', '就是', '这个', '那个'
])
FONT_PATH = None
if platform.system() == "Windows":
    FONT_PATHS = [
        'C:/Windows/Fonts/simhei.ttf',
        'C:/Windows/Fonts/msyh.ttc',
        'C:/Windows/Fonts/simsun.ttc'
    ]
    for path in FONT_PATHS:
        if os.path.exists(path):
            FONT_PATH = path
            break
elif platform.system() == "Darwin":
    FONT_PATHS = ['/System/Library/Fonts/PingFang.ttc', '/Library/Fonts/Arial Unicode.ttf']
    for path in FONT_PATHS:
        if os.path.exists(path):
            FONT_PATH = path
            break
else:  # Linux
    FONT_PATHS = ['/usr/share/fonts/truetype/wqy/wqy-microhei.ttc', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf']
    for path in FONT_PATHS:
        if os.path.exists(path):
            FONT_PATH = path
            break
WC_CONFIG = {
    'font_path': FONT_PATH,
    'width': 800,
    'height': 600,
    'background_color': 'white',
    'max_words': 200,
    'max_font_size': 150,
    'min_font_size': 10,
    'colormap': 'viridis',
    'random_state': 42,
}
def smart_path_process(file_path: str) -> str:
    """智能处理路径：去除引号、转换斜杠、自动补全后缀"""
    file_path = file_path.strip().strip('"\'')
    file_path = file_path.replace('\\', '/')
    if file_path.lower().endswith('.doc') and not file_path.lower().endswith('.docx'):
        new_path = file_path[:-4] + '.docx'
        print(f"⚠️  检测到.doc格式，自动转换路径为：{new_path}")
        print("   请确保已将.doc文件另存为.docx格式！")
        file_path = new_path
    return file_path
def validate_and_find_file(file_path: str) -> str:
    """验证文件，找不到时自动搜索当前目录"""
    file_path = smart_path_process(file_path)
    if os.path.exists(file_path) and os.path.isfile(file_path):
        return file_path
    current_dir = os.getcwd()
    docx_files = [f for f in os.listdir(current_dir) if f.lower().endswith('.docx')]
    if not docx_files:
        print(f"❌ 未找到文件：{file_path}")
        print(f"   当前目录（{current_dir}）下也没有找到任何.docx文件")
        return ""
    print(f"\n📂 当前目录下找到以下.docx文件：")
    for i, f in enumerate(docx_files, 1):
        print(f"   {i}. {f}")
    while True:
        choice = input(f"\n请选择文件序号（1-{len(docx_files)}）：").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(docx_files):
            selected = os.path.join(current_dir, docx_files[int(choice)-1])
            print(f"✅ 已选择：{selected}")
            return selected
        print("❌ 输入无效，请输入正确的序号")
def extract_word_text(file_path: str) -> str:
    """提取Word文本（增强容错）"""
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text.append(para_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text.append(cell_text)
        full_text = re.sub(r'\s+', ' ', ' '.join(text)).strip()
        full_text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s]', '', full_text)
        return full_text
    except Exception as e:
        print(f"❌ 读取文件失败：{str(e)}")
        return ""
def preprocess_text(text: str) -> list:
    """文本预处理"""
    if not text:
        return []
    words = jieba.lcut(text)
    filtered_words = [
        word for word in words
        if word not in STOPWORDS 
        and len(word) > 1 
        and not word.isdigit()
    ]
    return filtered_words
def generate_wordcloud(words: list, save_name: str = "词云结果.png"):
    """生成词云（自动处理掩码）"""
    if not words:
        print("❌ 无有效词语生成词云")
        return
    word_freq = Counter(words)
    try:
        wc = WordCloud(**WC_CONFIG)
        wc.generate_from_frequencies(word_freq)
        wc.to_file(save_name)
        print(f"\n✅ 词云已保存到：{os.path.abspath(save_name)}")
        plt.figure(figsize=(10, 8))
        plt.imshow(wc, interpolation='bilinear')
        plt.axis('off')
        plt.tight_layout()
        plt.show()
    except Exception as e:
        print(f"❌ 生成词云失败：{str(e)}"
def main():
    print("="*50)
    print("🎯 Word文档词云生成工具（全自动版）")
    print("="*50)
    print(f"📌 当前运行目录：{os.getcwd()}")
    print("💡 提示：直接按回车可自动选择当前目录下的docx文件\n")
    file_input = input("请输入Word文件路径（如：D:/python/wuren.docx）：").strip()
    if not file_input:
        file_input = "自动搜索"
    file_path = validate_and_find_file(file_input)
    if not file_path:
        return
    print("\n🔍 正在提取文本...")
    text = extract_word_text(file_path)
    if not text:
        print("❌ 未提取到任何文本")
        return
    print(f"✅ 文本提取完成，共 {len(text)} 个字符")
    print("\n🔧 正在处理文本（分词+去停用词）...")
    words = preprocess_text(text)
    if not words:
        print("❌ 处理后无有效词语")
        return
    print(f"✅ 文本处理完成，共 {len(words)} 个有效词语")
    top_10 = Counter(words).most_common(10)
    print("\n📊 高频词TOP10：")
    for i, (word, count) in enumerate(top_10, 1):
        print(f"   {i}. {word} - {count}次")
    save_name = input("\n请输入词云保存文件名（默认：词云结果.png）：").strip()
    if not save_name:
        save_name = "词云结果.png"
    if not save_name.lower().endswith(('.png', '.jpg', '.jpeg')):
        save_name += ".png"
    print("\n🎨 正在生成词云...")
    generate_wordcloud(words, save_name)
    print("\n🎉 操作完成！")
if __name__ == "__main__":
    main()
    
